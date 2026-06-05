"""Generate a sample Word template (.docx) showing how placeholders look."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ---- Header / company block ----
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("{{company_name}}")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x8C)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("{{company_address}}\n{{company_email}}  |  {{company_phone}}").font.size = Pt(9)

doc.add_paragraph("_" * 75)

# ---- Meta row ----
meta = doc.add_paragraph()
meta.add_run("Date: ").bold = True
meta.add_run("{{quote_date}}")
meta.add_run("\t\t\tQuote #: ").bold = True
meta.add_run("{{quote_id}}")

# ---- Recipient ----
doc.add_paragraph().add_run("To:").bold = True
to = doc.add_paragraph()
to.add_run("{{client_name}}\n{{client_company}}\n{{client_address}}")

# ---- Body ----
doc.add_paragraph()
g = doc.add_paragraph()
g.add_run("Dear {{client_name}},")
doc.add_paragraph(
    "Thank you for your interest in {{product_name}}. Please find below our "
    "quotation for {{project_description}}."
)

# ---- Items table (with docxtpl loop tags) ----
table = doc.add_table(rows=1, cols=4)
table.style = "Light Grid Accent 1"
hdr = table.rows[0].cells
for i, label in enumerate(["Item", "Qty", "Price", "Total"]):
    hdr[i].paragraphs[0].add_run(label).bold = True

# loop-start row
row_start = table.add_row().cells
row_start[0].text = "{%tr for i in items %}"

# the actual repeating data row
row = table.add_row().cells
row[0].text = "{{i.name}}"
row[1].text = "{{i.qty}}"
row[2].text = "{{i.rate}}"
row[3].text = "{{i.total}}"

# loop-end row
row_end = table.add_row().cells
row_end[0].text = "{%tr endfor %}"

# ---- Totals ----
doc.add_paragraph()
for label, ph in [("Subtotal:", "{{subtotal}}"), ("Tax:", "{{tax}}"),
                  ("Grand Total:", "{{grand_total}}")]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run(f"{label} ").bold = True
    p.add_run(ph)

doc.add_paragraph("This quote is valid until {{valid_until}}.")

# ---- Signature ----
doc.add_paragraph()
doc.add_paragraph("Sincerely,")
sig = doc.add_paragraph()
sig.add_run("{{sender_name}}\n{{sender_title}}")

doc.save("/home/user/temp2/sample_template.docx")
print("saved sample_template.docx")
